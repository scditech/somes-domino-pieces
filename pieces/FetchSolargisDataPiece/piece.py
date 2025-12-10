from domino.base_piece import BasePiece
from .models import InputModel, OutputModel
import pandas as pd
from pathlib import Path
from docx import Document

class FetchSolargisDataPiece(BasePiece):
    
    def _read_docx_files(self, file_paths):
        """Extract data from DOCX files."""
        csv_data = []
        
        for doc_file in file_paths:
            if not doc_file.exists():
                print(f"[WARNING] File not found: {doc_file}")
                continue
                
            print(f"[INFO] Processing DOCX document: {doc_file}")
            doc = Document(doc_file)
            csv_started = False

            if doc_file == file_paths[0]:
                # 1st doc file - we keep CSV header line                       
                # Process each paragraph in the document
                for paragraph in doc.paragraphs:
                    line = paragraph.text.strip()
                    # Check if we've reached the CSV data section
                    if "#Data:" in line:
                        csv_started = True
                        continue

                    # If we're in the CSV section, store the line
                    if csv_started and line:
                        csv_data.append(line)
            else:
                # next doc files - we don't keep CSV header line
                # Process each paragraph in the document
                for paragraph in doc.paragraphs:
                    line = paragraph.text.strip()
                    # Check if we've reached the CSV data section
                    if line.startswith("Date;"):
                        csv_started = True
                        continue
                    # If we're in the CSV section, store the line
                    if csv_started and line:
                        csv_data.append(line)
            
        return csv_data
    
    def _read_csv_files(self, file_paths):
        """Read and concatenate data from CSV files."""
        csv_data = []
        
        for csv_file in file_paths:
            if not csv_file.exists():
                print(f"[WARNING] File not found: {csv_file}")
                continue
                
            print(f"[INFO] Processing CSV file: {csv_file}")
            csv_started = False

            if csv_file == file_paths[0]:
                # 1st csv_file - we keep CSV header line
            
                # Read file line by line
                with open(csv_file, 'r') as f:
                    for line in f:
                        line = line.strip()
                        # Check if we've reached the CSV data section
                        if "#Data:" in line:
                            csv_started = True
                            continue
                        
                        # If we're in the CSV section, store the line
                        if csv_started and line:
                            csv_data.append(line)
            else:
                # next csv_files - we don't keep CSV header line
                # Read file line by line
                with open(csv_file, 'r') as f:
                    for line in f:
                        line = line.strip()
                        # Check if we've reached the CSV data section
                        if line.startswith("Date;"):
                            csv_started = True
                            continue
                        # If we're in the CSV section, store the line
                        if csv_started and line:
                            csv_data.append(line)
          
        return csv_data
    
    def _resolve_file_paths(self, input_path_str, file_type):
        """Resolve wildcard patterns and return list of matching files."""
        input_path = Path(input_path_str)
        file_extension = '.docx' if file_type.lower() == 'docx' else '.csv'
        
        # Resolve wildcard patterns in input path
        if '*' in str(input_path) or '?' in str(input_path):
            # Handle wildcard patterns using Path.glob
            files = [f for f in input_path.parent.glob(input_path.name) 
                     if f.suffix == file_extension]
        else:
            # Check if it's a directory
            if input_path.is_dir():
                # If directory, find all files of specified type
                pattern = f'*{file_extension}'
                files = list(input_path.glob(pattern))
            else:
                # Single file
                files = [input_path]
        
        return files
    
    def piece_function(self, input_data: InputModel):

        print(f"[INFO] Fetching data from {input_data.input_path}")
        print(f"[INFO] File type: {input_data.input_filetype}")

        csv_data = []
        
        # Resolve file paths based on wildcards
        file_paths = self._resolve_file_paths(input_data.input_path, input_data.input_filetype)
        
        if not file_paths:
            message = f"No {input_data.input_filetype.upper()} files found at: {input_data.input_path}"
            print(f"[WARNING] {message}")
            return OutputModel(
                message=message,
                file_path=""
            )
        
        # Process files based on type
        if input_data.input_filetype.lower() == 'docx':
            csv_data = self._read_docx_files(file_paths)
        elif input_data.input_filetype.lower() == 'csv':
            csv_data = self._read_csv_files(file_paths)
        else:
            message = f"Unsupported file type: {input_data.input_filetype}"
            print(f"[ERROR] {message}")
            return OutputModel(
                message=message,
                file_path=""
            )

        # Convert to DataFrame and write to CSV
        if csv_data:
            df = pd.DataFrame(csv_data)
            message = f"Files processed successfully, {len(df)} rows found."
            print(f"[SUCCESS] {message}")
            file_path = str(Path(self.results_path) / "raw_solargis.csv")
            df.to_csv(file_path, index=False, header=False)
            
        else:
            message = f"No data found in {input_data.input_filetype.upper()} files."
            print(f"[WARNING] {message}")
            file_path = None
            
        # Set display result
        self.display_result = {
            "file_type": "csv",
            "file_path": file_path if csv_data else ""
        }

        # Return output
        return OutputModel(
            message=message,
            file_path=file_path if csv_data else ""
        )
